"""
DOCX formatter — renders generated section content into a styled .docx
matching the reference document (IAR_3940_SG).

Style mapping:
  Heading 1  -> purple box, white text (Heading 1 New)
  Heading 2  -> dark navy bold
  Heading 3  -> dark navy bold, indented
  Bar Text   -> body paragraphs (2in left indent)
  Bar Text Bullets -> bulleted lists
  Bar Text - Important -> lavender callout boxes
  FD Question/Answer/Bottom -> Knowledge Check blocks
  CE LO Head -> course title (large purple, right-aligned)
"""

import logging
import os
import re
import tempfile
import hashlib
from pathlib import Path

from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

from ..constants.styles import (
    setup_styles,
    apply_heading1_shading,
    apply_heading2_accent,
    apply_important_shading,
    apply_fd_question_borders,
    BODY_FONT, HEADING_FONT, TITLE_FONT, QUOTE_FONT,
    BODY_SIZE, H4_SIZE, TITLE_SIZE,
    BODY_LEFT_INDENT, H4_LEFT_INDENT,
    DEEP_PURPLE, DARK_NAVY, NAVY_BLUE,
)

logger = logging.getLogger(__name__)
_DOCX_NATIVE_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"}


def _coerce_docx_compatible_image_path(path_str: str, *, force_convert: bool = False) -> str:
    """Convert non-native image formats into a PNG sidecar that python-docx can embed."""
    path = Path(path_str)
    if not force_convert and path.suffix.lower() in _DOCX_NATIVE_IMAGE_EXTS:
        return str(path)

    cache_dir = Path(tempfile.gettempdir()) / "lectora_docx_image_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    stat = path.stat()
    cache_key = hashlib.sha256(
        f"{path.resolve()}:{stat.st_mtime_ns}:{force_convert}".encode("utf-8")
    ).hexdigest()[:16]
    converted = cache_dir / f"{path.stem}_{cache_key}.png"
    if converted.exists():
        return str(converted)

    with Image.open(path) as src:
        normalized = (
            src.convert("RGBA")
            if src.mode in ("P", "LA", "RGBA")
            else src.convert("RGB")
        )
        normalized.save(converted, format="PNG")

    logger.info(
        "[A2] Converted image for DOCX embedding: %s -> %s",
        path.name,
        converted.name,
    )
    return str(converted)


# ---------------------------------------------------------------------------
# Bookmark & hyperlink helpers (clickable TOC)
# ---------------------------------------------------------------------------

def _make_bookmark_name(heading: str, bm_id: int) -> str:
    """Create a valid Word bookmark name from a heading string."""
    clean = re.sub(r"[^a-zA-Z0-9]", "_", (heading or "").lower())
    clean = re.sub(r"_+", "_", clean).strip("_")[:25]
    return f"s{bm_id}_{clean}"


def _apply_bookmark(para, bm_id: int, bm_name: str) -> None:
    """Insert w:bookmarkStart / w:bookmarkEnd around a paragraph's content."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bm_id))
    bm_start.set(qn("w:name"), bm_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bm_id))

    para._p.insert(0, bm_start)
    para._p.append(bm_end)


def _add_toc_hyperlink_para(
    doc,
    text: str,
    anchor: str,
    color: RGBColor,
    font_size_pt: int = 11,
    indent_inches: float | None = None,
) -> None:
    """Add a TOC paragraph whose text is a clickable internal hyperlink (w:anchor)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    if indent_inches:
        para.paragraph_format.left_indent = Inches(indent_inches)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)

    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), str(color))  # RGBColor.__str__ returns "RRGGBB" hex
    rPr.append(color_elem)

    for tag in ("w:sz", "w:szCs"):
        sz = OxmlElement(tag)
        sz.set(qn("w:val"), str(font_size_pt * 2))  # half-points
        rPr.append(sz)

    run_elem.append(rPr)

    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run_elem.append(t_elem)

    hyperlink.append(run_elem)
    para._p.append(hyperlink)


# ---------------------------------------------------------------------------

_RESERVED_HEADING_RE = re.compile(
    r"^\s*(\d+(\.\d+)*\s+)?"
    r"(overview|learning\s+objectives?|learning\s+outcomes?|course\s+objectives?|"
    r"summary|assessment|introduction)\s*$",
    re.IGNORECASE,
)


def _inject_missing_lesson_parent_sections(sections: list[dict]) -> list[dict]:
    """Ensure each lesson title is rendered before its generated subtopic sections.

    A2 stores the lesson title in ``outline_lesson`` for every generated subtopic,
    but when no parent overview block is generated the final DOCX would otherwise
    render only the subtopic headings. This helper inserts a synthetic level-1
    heading once per lesson so the document structure matches the TO outline:

        3.0 Lesson Title
        3.1 First subtopic
        3.2 Second subtopic

    The injected parent section includes a bullet list of its subtopic headings so
    both the DOCX body and the course editor have visible content for every lesson
    group, not just a bare heading box.
    """
    # Pass 1: collect subtopic headings per outline_lesson so synthetic parent
    # sections can include a brief "Topics in this section" overview.
    lesson_subtopic_headings: dict[str, list[str]] = {}
    for sec in sections:
        ol = (sec.get("outline_lesson") or "").strip()
        lvl = sec.get("level", 2)
        hd = (sec.get("heading") or "").strip()
        is_parent = bool(
            lvl == 1
            or sec.get("is_parent_overview")
            or (ol and hd == ol)
        )
        if ol and not is_parent and hd:
            lesson_subtopic_headings.setdefault(ol, []).append(hd)

    # Pass 2: inject synthetic L1 parents where the lesson has no existing parent.
    result: list[dict] = []
    current_outline_lesson = ""
    lesson_has_parent = False

    for sec in sections:
        sec = dict(sec)
        outline_lesson = (sec.get("outline_lesson") or "").strip()
        heading = (sec.get("heading") or "").strip()
        level = sec.get("level", 2)

        if outline_lesson != current_outline_lesson:
            current_outline_lesson = outline_lesson
            lesson_has_parent = False

        is_existing_parent = bool(
            level == 1
            or sec.get("is_parent_overview")
            or (outline_lesson and heading == outline_lesson)
        )

        if outline_lesson and not lesson_has_parent and not is_existing_parent:
            subtopic_titles = lesson_subtopic_headings.get(outline_lesson, [])
            body_paragraphs: list[dict] = []
            if subtopic_titles:
                body_paragraphs = [{"type": "bullet_list", "items": subtopic_titles}]
            result.append(
                {
                    "heading": outline_lesson,
                    "level": 1,
                    "status": "skipped_thin",
                    "body_paragraphs": body_paragraphs,
                    "word_count": 0,
                    "outline_lesson": outline_lesson,
                    "images": [],
                    "subtopics": subtopic_titles,
                    "maps_to_objectives": [],
                    "section_id": "",
                    "attempts": 1,
                }
            )
            lesson_has_parent = True

        if is_existing_parent:
            lesson_has_parent = True

        result.append(sec)

    return result


def _renumber_sections(sections: list[dict]) -> list[dict]:
    """Re-number generated sections sequentially starting at 3.0.

    Fixed sections 1.0 OVERVIEW and 2.0 Learning Objectives are rendered by
    doc_formatter from metadata and must NEVER appear in this list.  If a
    reserved heading does slip through (e.g. from a malformed TO), it is
    removed here so it never occupies a content-section slot or consumes a
    major-number counter tick.

    Level-1 sections (parent overviews) get N.0 numbers (3.0, 4.0, …).
    Level-2 sections (subtopics)        get N.M numbers (3.1, 3.2, …).
    Guarantees proper ordering: a level-2 section always follows its level-1 parent.
    """
    _NUM_PREFIX = re.compile(r"^\d+(\.\d+)*[\s.:\-]*")

    major = 2   # 1.0 OVERVIEW + 2.0 LOs are fixed; first content section → 3.0
    minor = 0

    result: list[dict] = []
    for sec in sections:
        sec = dict(sec)
        level = sec.get("level", 2)
        heading = (sec.get("heading") or "").strip()
        clean = _NUM_PREFIX.sub("", heading).strip() or heading

        # Reserved headings belong in front-matter (rendered from metadata),
        # not in the generated content body.  Drop them silently.
        if _RESERVED_HEADING_RE.match(clean):
            continue

        if level == 1:
            major += 1
            minor = 0
            sec["heading"] = f"{major}.0 {clean}"
        elif level == 2:
            minor += 1
            sec["heading"] = f"{major}.{minor} {clean}"
        # level 3+ (heading_3 body-paragraphs) keep as-is

        result.append(sec)
    return result


def _add_bold_run(paragraph, text: str, font_name: str = None,
                  font_size=None, color: RGBColor = None):
    """Add a bold run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = True
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if color:
        run.font.color.rgb = color
    return run


def _apply_body_indent(paragraph):
    """Apply 2in left indent to a body paragraph (Bar Text style)."""
    paragraph.paragraph_format.left_indent = BODY_LEFT_INDENT
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)


def _render_text_with_bold(paragraph, text: str, font_name: str = None,
                           font_size=None, color: RGBColor = None):
    """
    Render text into a paragraph, converting **bold** markers to actual bold runs.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if color:
            run.font.color.rgb = color


def _add_title_page(doc, course_title: str):
    """Add course title in CE LO Head style (large purple, right-aligned)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(course_title)
    run.font.name = TITLE_FONT
    run.font.size = TITLE_SIZE
    run.font.color.rgb = DEEP_PURPLE


def _add_toc(
    doc,
    sections: list[dict],
    *,
    bookmark_map: dict[str, tuple[int, str]] | None = None,
    conclusion_heading: str | None = None,
    lo_heading: str = "2.0 Learning Objectives",
    include_overview: bool = True,
) -> None:
    """Add a Table of Contents with clickable internal hyperlinks."""
    bm = bookmark_map or {}

    def _entry(text: str, indent: float | None = None) -> None:
        bm_data = bm.get(text)
        if bm_data:
            _add_toc_hyperlink_para(doc, text, bm_data[1], DARK_NAVY, 11, indent_inches=indent)
        else:
            tp = doc.add_paragraph()
            run = tp.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_NAVY
            if indent:
                tp.paragraph_format.left_indent = Inches(indent)
            tp.paragraph_format.space_before = Pt(2)
            tp.paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_before = Pt(12)
    if include_overview:
        _entry("1.0 OVERVIEW")
    _entry(lo_heading)

    for sec in sections:
        level = sec.get("level", 1)
        heading = sec.get("heading", "")
        if not heading:
            continue
        _entry(heading, indent=0.25 if level >= 2 else None)

    if conclusion_heading:
        _entry(conclusion_heading)


def _add_section_1_overview(
    doc,
    description: str,
    bookmark_map: dict[str, tuple[int, str]] | None = None,
) -> None:
    """Section 1.0 — OVERVIEW (description only; LOs are section 2.0)."""
    h1 = doc.add_heading("1.0 OVERVIEW", level=1)
    apply_heading1_shading(h1)
    bm_data = (bookmark_map or {}).get("1.0 OVERVIEW")
    if bm_data:
        _apply_bookmark(h1, *bm_data)

    doc.add_heading("Description", level=2)
    desc_stripped = (description or "").strip()
    if desc_stripped:
        blocks = [b.strip() for b in desc_stripped.split("\n\n") if b.strip()]
        if not blocks:
            blocks = [desc_stripped]
        for block in blocks:
            p = doc.add_paragraph()
            _apply_body_indent(p)
            _render_text_with_bold(p, block, font_name=BODY_FONT, font_size=BODY_SIZE)


def _infer_conclusion_heading(generated_sections: list[dict]) -> str:
    """Numbered conclusion after last major (N.0) content section, else plain 'Conclusion'."""
    max_major = 2
    for sec in generated_sections:
        h = (sec.get("heading") or "").strip()
        m = re.match(r"^(\d+)\.0\b", h)
        if m:
            max_major = max(max_major, int(m.group(1)))
    if max_major > 2:
        return f"{max_major + 1}.0 Conclusion"
    return "Conclusion"


def _add_section_2_learning_objectives(
    doc,
    learning_objectives: list[str],
    bookmark_map: dict[str, tuple[int, str]] | None = None,
) -> None:
    """Section 2.0 — Learning Objectives (course-level list)."""
    los_for_doc = [str(lo).strip() for lo in (learning_objectives or []) if lo and str(lo).strip()]
    lo_heading = (
        "2.0 Learning Objectives"
        if len(los_for_doc) != 1
        else "2.0 Learning Objective"
    )
    h1 = doc.add_heading(lo_heading, level=1)
    apply_heading1_shading(h1)
    bm_data = (bookmark_map or {}).get(lo_heading)
    if bm_data:
        _apply_bookmark(h1, *bm_data)

    for lo in los_for_doc:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent = BODY_LEFT_INDENT
        _render_text_with_bold(bp, lo, font_name=BODY_FONT, font_size=BODY_SIZE)


def _add_conclusion_section(
    doc,
    conclusion_text: str,
    generated_sections: list[dict],
    bookmark_map: dict[str, tuple[int, str]] | None = None,
) -> None:
    """Final section — recap for students (numbered after last major chapter when possible)."""
    heading = _infer_conclusion_heading(generated_sections)
    h1 = doc.add_heading(heading, level=1)
    apply_heading1_shading(h1)
    bm_data = (bookmark_map or {}).get(heading)
    if bm_data:
        _apply_bookmark(h1, *bm_data)

    text = (conclusion_text or "").strip()
    if not text:
        p = doc.add_paragraph()
        _apply_body_indent(p)
        run = p.add_run("[Conclusion — generation pending or empty]")
        run.italic = True
        return

    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    if not blocks:
        blocks = [text]
    for block in blocks:
        p = doc.add_paragraph()
        _apply_body_indent(p)
        _render_text_with_bold(p, block, font_name=BODY_FONT, font_size=BODY_SIZE)


def _render_table(doc, table_data: dict) -> None:
    """Render a structured table (comparison matrix / process table) into the document."""
    headers = table_data.get("headers") or []
    rows = table_data.get("rows") or []
    caption = (table_data.get("caption") or "").strip()

    n_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    if n_cols == 0:
        return

    # Optional caption above the table
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.left_indent = BODY_LEFT_INDENT
        cap_p.paragraph_format.space_before = Pt(8)
        cap_p.paragraph_format.space_after = Pt(2)
        cap_run = cap_p.add_run(caption)
        cap_run.bold = True
        cap_run.font.name = BODY_FONT
        cap_run.font.size = BODY_SIZE
        cap_run.font.color.rgb = DARK_NAVY

    n_rows = (1 if headers else 0) + len(rows)
    if n_rows == 0:
        return

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    row_offset = 0
    if headers:
        hdr_row = table.rows[0]
        for i, hdr_text in enumerate(headers[:n_cols]):
            cell = hdr_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(hdr_text))
            run.bold = True
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE
            # Light purple header shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "E8E0F0")
            tcPr.append(shd)
        row_offset = 1

    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + row_offset]
        for c_i, cell_text in enumerate(list(row_data)[:n_cols]):
            cell = row.cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            _render_text_with_bold(p, str(cell_text), font_name=BODY_FONT, font_size=BODY_SIZE)

    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(2)

def _insert_image(doc, img: dict, max_width_inches: float = 4.5):
    """
    Insert a single image into the document from its saved_path.

    Respects original aspect ratio; caps width at max_width_inches.
    Adds caption below if present (from doc text only — never AI-generated).
    """
    def _add_placeholder(reason: str) -> None:
        p = doc.add_paragraph()
        _apply_body_indent(p)
        run = p.add_run(f"[Image: {img.get('media_filename', 'missing')} — {reason}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    saved_path = img.get("saved_path", "")
    if not saved_path or not os.path.isfile(saved_path):
        # Image file missing — add placeholder
        _add_placeholder("file not found")
        return

    # Determine display width from original size or default
    size_cm = img.get("size_cm", {})
    orig_w = size_cm.get("width")
    orig_h = size_cm.get("height")

    if orig_w and orig_w > 0:
        # Convert cm to inches; cap at max_width_inches
        width_in = min(orig_w / 2.54, max_width_inches)
    else:
        width_in = max_width_inches

    # Image paragraph — centered on page; no left indent (would shift image right-of-centre)
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(8)
    img_para.paragraph_format.space_after = Pt(4)

    picture_path = saved_path
    try:
        picture_path = _coerce_docx_compatible_image_path(saved_path)
        run = img_para.add_run()
        run.add_picture(picture_path, width=Inches(width_in))
    except UnrecognizedImageError:
        try:
            picture_path = _coerce_docx_compatible_image_path(
                saved_path,
                force_convert=True,
            )
            run = img_para.add_run()
            run.add_picture(picture_path, width=Inches(width_in))
        except Exception as exc:
            logger.warning("[A2] Could not render image %s: %s", saved_path, exc)
            _add_placeholder("unsupported image format")
            return
    except Exception as exc:
        logger.warning("[A2] Could not render image %s: %s", saved_path, exc)
        _add_placeholder("could not be rendered")
        return

    # Caption (only from doc text — no AI-generated descriptions)
    caption = img.get("caption", "")
    alt_text = img.get("alt_text", "")
    caption_text = caption or alt_text

    if caption_text:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.left_indent = BODY_LEFT_INDENT
        cap_para.paragraph_format.space_before = Pt(0)
        cap_para.paragraph_format.space_after = Pt(8)
        cap_run = cap_para.add_run(caption_text)
        cap_run.italic = True
        cap_run.font.name = BODY_FONT
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _insert_section_images(doc, images: list[dict]):
    """
    Insert all images mapped to a section.
    Called after rendering body paragraphs so images appear at end of section,
    matching the reference doc pattern where images follow their context.
    """
    if not images:
        return

    for img in images:
        _insert_image(doc, img)


def _render_section_content(
    doc,
    section: dict,
    bookmark_map: dict[str, tuple[int, str]] | None = None,
) -> None:
    """Render a single section's body_paragraphs and images into the document."""
    heading = section.get("heading", "")
    level = section.get("level", 2)
    bm = bookmark_map or {}

    # Add heading with bookmark
    if level == 1:
        h = doc.add_heading(heading, level=1)
        apply_heading1_shading(h)
    elif level == 2:
        h = doc.add_heading(heading, level=2)
        apply_heading2_accent(h)
    elif level == 3:
        h = doc.add_heading(heading, level=3)
    else:
        h = doc.add_heading(heading, level=2)
        apply_heading2_accent(h)

    bm_data = bm.get(heading)
    if bm_data:
        _apply_bookmark(h, *bm_data)

    # Render body paragraphs
    for para in section.get("body_paragraphs", []):
        ptype = para.get("type", "text")

        if ptype == "text":
            p = doc.add_paragraph()
            _apply_body_indent(p)
            _render_text_with_bold(
                p, para.get("content", ""),
                font_name=BODY_FONT, font_size=BODY_SIZE,
            )

        elif ptype == "heading_3":
            doc.add_heading(para.get("content", ""), level=3)

        elif ptype == "heading_4":
            h4 = doc.add_paragraph()
            h4.paragraph_format.left_indent = H4_LEFT_INDENT
            h4.paragraph_format.space_before = Pt(8)
            _add_bold_run(
                h4, para.get("content", ""),
                font_name=HEADING_FONT, font_size=H4_SIZE, color=DARK_NAVY,
            )

        elif ptype == "bullet_list":
            for item in para.get("items", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent = BODY_LEFT_INDENT
                _render_text_with_bold(
                    bp, item, font_name=BODY_FONT, font_size=BODY_SIZE,
                )

        elif ptype == "sub_bullet_list":
            for item in para.get("items", []):
                bp = doc.add_paragraph(style="List Bullet 2")
                bp.paragraph_format.left_indent = Inches(2.5)
                _render_text_with_bold(
                    bp, item, font_name=BODY_FONT, font_size=BODY_SIZE,
                )

        elif ptype == "numbered_list":
            for item in para.get("items", []):
                np = doc.add_paragraph(style="List Number")
                np.paragraph_format.left_indent = BODY_LEFT_INDENT
                _render_text_with_bold(
                    np, item, font_name=BODY_FONT, font_size=BODY_SIZE,
                )

        elif ptype == "important_callout":
            imp = doc.add_paragraph()
            imp.paragraph_format.left_indent = BODY_LEFT_INDENT
            imp.paragraph_format.space_before = Pt(6)
            imp.paragraph_format.space_after = Pt(6)
            apply_important_shading(imp)
            label = (para.get("label") or "").strip()
            body = (para.get("content") or "").strip()
            callout_text = f"**{label}:** {body}" if label else body
            _render_text_with_bold(
                imp, callout_text,
                font_name=BODY_FONT, font_size=BODY_SIZE, color=NAVY_BLUE,
            )

        elif ptype == "table":
            _render_table(doc, para)

    # Insert images mapped to this section (from A0 extraction -> A1 mapping)
    _insert_section_images(doc, section.get("images", []))


def build_study_guide_docx(
    course_title: str,
    course_description: str,
    learning_objectives: list[str],
    generated_sections: list[dict],
    output_path: str,
    *,
    conclusion_text: str = "",
    include_overview: bool = True,
) -> str:
    """
    Assemble the full study guide .docx from generated sections.

    Follows the reference doc style:
      1. Title page (CE LO Head)
      2. Table of Contents
      3. 1.0 OVERVIEW (description)
      4. 2.0 Learning Objectives
      5. Section content from 3.0 onward (headings + body + KCs)
      6. Conclusion (final recap)

    Returns the output file path.
    """
    doc = Document()
    setup_styles(doc)

    # Ensure each outline lesson title appears before its generated subtopic sections.
    generated_sections = _inject_missing_lesson_parent_sections(generated_sections)

    # Re-number sections: 3.0 → 3.1 → 3.2 → 4.0 → 4.1 … (offsets fixed 1.0/2.0)
    generated_sections = _renumber_sections(generated_sections)

    conclusion_heading = _infer_conclusion_heading(generated_sections)

    # Compute LO heading (singular vs plural) for consistent TOC ↔ content anchor
    los_for_doc = [str(lo).strip() for lo in (learning_objectives or []) if lo and str(lo).strip()]
    lo_heading = "2.0 Learning Objectives" if len(los_for_doc) != 1 else "2.0 Learning Objective"

    # Build bookmark map: every heading that will appear in the document
    _all_headings: list[str] = []
    if include_overview:
        _all_headings.append("1.0 OVERVIEW")
    _all_headings.append(lo_heading)
    for sec in generated_sections:
        h = (sec.get("heading") or "").strip()
        if h:
            _all_headings.append(h)
    _all_headings.append(conclusion_heading)

    bookmark_map: dict[str, tuple[int, str]] = {}
    for _bm_id, _heading in enumerate(dict.fromkeys(h for h in _all_headings if h)):
        bookmark_map[_heading] = (_bm_id, _make_bookmark_name(_heading, _bm_id))

    # 1. Title page
    _add_title_page(doc, course_title)
    doc.add_page_break()

    # 2. Table of Contents (clickable)
    _add_toc(
        doc,
        generated_sections,
        bookmark_map=bookmark_map,
        conclusion_heading=conclusion_heading,
        lo_heading=lo_heading,
        include_overview=include_overview,
    )
    doc.add_page_break()

    if include_overview:
        _add_section_1_overview(doc, course_description, bookmark_map=bookmark_map)
    _add_section_2_learning_objectives(doc, learning_objectives, bookmark_map=bookmark_map)

    # 5. Main course content (typically numbered 3.0+ in section headings)
    for section in generated_sections:
        status = section.get("status", "")
        if status == "skipped_thin":
            level = section.get("level", 1)
            heading = section.get("heading", "")
            if heading:
                if level == 1:
                    h = doc.add_heading(heading, level=1)
                    apply_heading1_shading(h)
                elif level == 2:
                    h = doc.add_heading(heading, level=2)
                    apply_heading2_accent(h)
                else:
                    h = doc.add_heading(heading, level=3)
                bm_data = bookmark_map.get(heading)
                if bm_data:
                    _apply_bookmark(h, *bm_data)
            # Render any body_paragraphs (e.g. the subtopic overview list injected
            # by _inject_missing_lesson_parent_sections so lesson groups have visible
            # content in the document body, not just a bare heading box).
            for para in section.get("body_paragraphs") or []:
                ptype = para.get("type", "text")
                if ptype == "bullet_list":
                    for item in (para.get("items") or []):
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.paragraph_format.left_indent = BODY_LEFT_INDENT
                        _render_text_with_bold(
                            bp, item, font_name=BODY_FONT, font_size=BODY_SIZE,
                        )
                elif ptype == "text":
                    p = doc.add_paragraph()
                    _apply_body_indent(p)
                    _render_text_with_bold(
                        p, para.get("content", ""),
                        font_name=BODY_FONT, font_size=BODY_SIZE,
                    )
            continue

        if status == "failed":
            heading = section.get("heading", "")
            h = doc.add_heading(heading, level=section.get("level", 2))
            bm_data = bookmark_map.get(heading)
            if bm_data:
                _apply_bookmark(h, *bm_data)
            p = doc.add_paragraph()
            _apply_body_indent(p)
            run = p.add_run("[Content generation failed — manual review required]")
            run.italic = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            continue

        _render_section_content(doc, section, bookmark_map=bookmark_map)

    # 6. Conclusion (end of course)
    _add_conclusion_section(doc, conclusion_text, generated_sections, bookmark_map=bookmark_map)

    # Save
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return str(output)
