"""
Document chunker for multi-file course generation.

Splits DOCX and PDF documents into semantic chunks suitable for
topic-wise retrieval in the content generation pipeline.

Architecture is designed to scale to 100+ files and supports future
vector/RAG integration.

Chunk schema:
  {
    "id":          str,          # "{source_stem}_{chunk_idx:04d}"
    "source":      str,          # original filename / label
    "heading":     str | None,   # nearest heading above this chunk
    "heading_level": int | None, # 1 = top-level, 2 = sub, etc.
    "content":     str,          # chunk text
    "word_count":  int,
    "para_start":  int | None,   # para_idx in source doc (DOCX only)
    "para_end":    int | None,
  }
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

# Default chunk sizes
_DEFAULT_CHUNK_WORDS = 400       # target words per chunk
_DEFAULT_OVERLAP_WORDS = 50      # overlap between adjacent chunks

# Compiled once at module level — used by both DOCX and PDF chunking paths
_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)[\s.\-:]\s*(.+)$")
_NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)[\s.\-:]\s*(.+)$")
_TOPIC_TOKEN_RE = re.compile(r"[a-z]{3,}")


# ---------------------------------------------------------------------------
# Module-level helper — DOCX heading detection
# ---------------------------------------------------------------------------

def _detect_heading(para: object) -> tuple[bool, int]:
    """Return (is_heading, level) for a python-docx paragraph.

    Recognises both Word built-in Heading styles and numbered prefixes
    (e.g. ``3.1 Title``).  Level 1 is returned when the style digit is
    absent or cannot be parsed.
    """
    style: str = para.style.name  # type: ignore[attr-defined]
    if "Heading" in style:
        try:
            level = int(style[-1]) if style[-1].isdigit() else 1
        except (IndexError, ValueError):
            level = 1
        return True, level
    m = _HEADING_RE.match(para.text.strip())  # type: ignore[attr-defined]
    if m:
        dots = m.group(1).count(".")
        return True, dots + 1
    return False, 0


# ---------------------------------------------------------------------------
# DocumentChunker — configurable chunker for DOCX and PDF sources
# ---------------------------------------------------------------------------

class DocumentChunker:
    """Splits DOCX and PDF source documents into overlapping semantic chunks.

    Configuration (chunk size, overlap) is set once at construction and
    applied to every chunking call on this instance.

    Example::

        chunker = DocumentChunker(chunk_words=400, overlap_words=50)
        chunks  = chunker.chunk_files(["study_guide.docx", "reference.pdf"])
    """

    def __init__(
        self,
        chunk_words: int = _DEFAULT_CHUNK_WORDS,
        overlap_words: int = _DEFAULT_OVERLAP_WORDS,
    ) -> None:
        self._chunk_words = chunk_words
        self._overlap_words = overlap_words

    # ── Public methods ────────────────────────────────────────────────────────

    def chunk_docx(self, docx_path: str, source_label: str | None = None) -> list[dict]:
        """Split a DOCX file into semantic chunks at heading boundaries.

        Within a heading section, further splits by word count so no single
        chunk exceeds ``chunk_words``.  Overlap preserves context across chunk
        boundaries.
        """
        try:
            from docx import Document
        except ImportError:
            logger.warning("[chunker] python-docx not available — skipping DOCX chunking.")
            return []

        path = Path(docx_path)
        label = source_label or path.stem
        doc = Document(str(path))

        chunk_words = self._chunk_words
        overlap_words = self._overlap_words

        chunks: list[dict] = []
        current_heading: str | None = None
        current_level: int | None = None
        buffer: list[str] = []
        buffer_para_start = 0
        chunk_idx = 0

        def _flush(para_end: int) -> None:
            nonlocal chunk_idx
            if not buffer:
                return
            text = " ".join(buffer)
            chunks.append({
                "id": f"{label}_{chunk_idx:04d}",
                "source": label,
                "heading": current_heading,
                "heading_level": current_level,
                "content": text,
                "word_count": len(text.split()),
                "para_start": buffer_para_start,
                "para_end": para_end,
            })
            chunk_idx += 1

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            is_head, head_level = _detect_heading(para)

            if is_head:
                _flush(max(0, idx - 1))
                buffer.clear()
                buffer_para_start = idx + 1
                current_heading = text
                current_level = head_level
                continue

            buffer.extend(text.split())

            while len(buffer) >= chunk_words + overlap_words:
                chunk_text = " ".join(buffer[:chunk_words])
                chunks.append({
                    "id": f"{label}_{chunk_idx:04d}",
                    "source": label,
                    "heading": current_heading,
                    "heading_level": current_level,
                    "content": chunk_text,
                    "word_count": chunk_words,
                    "para_start": buffer_para_start,
                    "para_end": idx,
                })
                chunk_idx += 1
                buffer = buffer[chunk_words - overlap_words:]
                buffer_para_start = idx

        _flush(len(doc.paragraphs) - 1)
        return chunks

    def chunk_pdf(self, pdf_path: str, source_label: str | None = None) -> list[dict]:
        """Split a PDF file into semantic chunks.

        Splits at numbered-heading boundaries (e.g. ``3.0 Topic``) when
        present, otherwise falls back to fixed-size word-window chunking.
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("[chunker] pypdf not available — skipping PDF chunking.")
            return []

        path = Path(pdf_path)
        label = source_label or path.stem

        try:
            reader = PdfReader(str(path))
        except Exception as exc:
            logger.warning("[chunker] Cannot read PDF %s: %s", pdf_path, exc)
            return []

        chunk_words = self._chunk_words
        overlap_words = self._overlap_words

        all_lines: list[str] = []
        for page in reader.pages:
            text = (page.extract_text() or "").strip()
            if text:
                all_lines.extend(text.splitlines())

        chunks: list[dict] = []
        chunk_idx = 0
        current_heading: str | None = None
        current_level: int | None = None
        buffer: list[str] = []

        def _flush() -> None:
            nonlocal chunk_idx
            if not buffer:
                return
            text = " ".join(buffer)
            chunks.append({
                "id": f"{label}_{chunk_idx:04d}",
                "source": label,
                "heading": current_heading,
                "heading_level": current_level,
                "content": text,
                "word_count": len(buffer),
                "para_start": None,
                "para_end": None,
            })
            chunk_idx += 1

        for raw_line in all_lines:
            line = raw_line.strip()
            if not line:
                continue

            m = _NUMBERED_HEADING_RE.match(line)
            if m and len(line) <= 120:
                _flush()
                buffer = []
                current_heading = line
                current_level = m.group(1).count(".") + 1
                continue

            buffer.extend(line.split())

            while len(buffer) >= chunk_words + overlap_words:
                chunk_text = " ".join(buffer[:chunk_words])
                chunks.append({
                    "id": f"{label}_{chunk_idx:04d}",
                    "source": label,
                    "heading": current_heading,
                    "heading_level": current_level,
                    "content": chunk_text,
                    "word_count": chunk_words,
                    "para_start": None,
                    "para_end": None,
                })
                chunk_idx += 1
                buffer = buffer[chunk_words - overlap_words:]

        _flush()
        return chunks

    def chunk_files(self, file_paths: list[str]) -> list[dict]:
        """Chunk all files (DOCX + PDF) and return a unified chunk list.

        Each chunk retains its ``source`` label so retrieval can cite origins.
        """
        all_chunks: list[dict] = []
        for path in file_paths:
            ext = Path(path).suffix.lower()
            label = Path(path).stem
            if ext == ".docx":
                file_chunks = self.chunk_docx(path, source_label=label)
            elif ext == ".pdf":
                file_chunks = self.chunk_pdf(path, source_label=label)
            else:
                logger.warning("[chunker] Unsupported file type: %s — skipping.", path)
                continue
            all_chunks.extend(file_chunks)
            logger.info("[chunker] %s → %d chunks", label, len(file_chunks))
        return all_chunks


# ---------------------------------------------------------------------------
# Retrieval utilities — operate on chunk lists, not source files
# ---------------------------------------------------------------------------

def retrieve_chunks_for_topic(
    topic: str,
    chunks: list[dict],
    top_k: int = 6,
    min_score: float = 0.05,
) -> list[dict]:
    """Simple keyword-overlap retrieval — returns top_k chunks most relevant to topic.

    This is a lightweight BM25-inspired scorer using term frequency overlap.
    Suitable for 100+ file scenarios without a vector database.

    Args:
        topic:     The TO section title / topic string.
        chunks:    Flat list of chunk dicts from ``chunk_files()``.
        top_k:     Maximum number of chunks to return.
        min_score: Minimum relevance score (0–1) to include a chunk.

    Returns:
        Ranked list of chunk dicts, most relevant first.
    """
    topic_tokens = set(_TOPIC_TOKEN_RE.findall(topic.lower()))
    if not topic_tokens:
        return chunks[:top_k]

    scored: list[tuple[float, dict]] = []
    for chunk in chunks:
        combined = ((chunk.get("heading") or "") + " " + chunk["content"]).lower()
        chunk_tokens = _TOPIC_TOKEN_RE.findall(combined)
        if not chunk_tokens:
            continue
        matches = sum(1 for t in chunk_tokens if t in topic_tokens)
        score = matches / len(chunk_tokens)
        if score >= min_score:
            scored.append((score, chunk))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [c for _, c in scored[:top_k]]


def build_context_for_topic(
    topic: str,
    chunks: list[dict],
    top_k: int = 6,
    max_words: int = 3000,
) -> str:
    """Retrieve relevant chunks for a topic and format them for an LLM prompt.

    Returns a formatted string of the most relevant source passages.
    This is the "smart retrieval" layer that feeds only relevant content to A2.
    """
    relevant = retrieve_chunks_for_topic(topic, chunks, top_k=top_k)
    if not relevant:
        return ""

    parts: list[str] = []
    total_words = 0
    for chunk in relevant:
        heading_label = f"[From: {chunk['source']}]"
        if chunk.get("heading"):
            heading_label += f" [{chunk['heading']}]"
        text = chunk["content"]
        words = text.split()
        if total_words + len(words) > max_words:
            remaining = max_words - total_words
            text = " ".join(words[:remaining]) + " […]"
        parts.append(f"{heading_label}\n{text}")
        total_words += len(words)
        if total_words >= max_words:
            break

    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# Backward-compatible module-level wrappers
# ---------------------------------------------------------------------------

def chunk_docx(
    docx_path: str,
    chunk_words: int = _DEFAULT_CHUNK_WORDS,
    overlap_words: int = _DEFAULT_OVERLAP_WORDS,
    source_label: str | None = None,
) -> list[dict]:
    return DocumentChunker(chunk_words, overlap_words).chunk_docx(docx_path, source_label=source_label)


def chunk_pdf(
    pdf_path: str,
    chunk_words: int = _DEFAULT_CHUNK_WORDS,
    overlap_words: int = _DEFAULT_OVERLAP_WORDS,
    source_label: str | None = None,
) -> list[dict]:
    return DocumentChunker(chunk_words, overlap_words).chunk_pdf(pdf_path, source_label=source_label)


def chunk_files(
    file_paths: list[str],
    chunk_words: int = _DEFAULT_CHUNK_WORDS,
    overlap_words: int = _DEFAULT_OVERLAP_WORDS,
) -> list[dict]:
    return DocumentChunker(chunk_words, overlap_words).chunk_files(file_paths)
