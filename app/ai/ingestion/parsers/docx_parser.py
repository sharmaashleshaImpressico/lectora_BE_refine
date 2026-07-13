from __future__ import annotations
import logging
import uuid

from app.ai.ingestion.chunking.models import BlockType, DocumentNode
from app.ai.ingestion.parsers.base import BaseDocumentParser

logger = logging.getLogger(__name__)

_A0_DOC_PARSER = (
    "app.ai.agents.to_generation_pipeline.step_01_parse_and_generate_outline"
    ".parse_documents.utils.doc_parser"
)


class DOCXParser(BaseDocumentParser):
    """
    Parse a DOCX into DocumentNode objects.

    Heading positions and levels come from CourseDocParser.extract_heading_tree()
    (A0's battle-tested logic), so numbered headings, Title/Subtitle styles, and
    all edge cases are handled consistently with the rest of the pipeline.
    """

    def parse(self, path: str) -> list[DocumentNode]:
        from importlib import import_module

        import docx as python_docx

        CourseDocParser = import_module(_A0_DOC_PARSER).CourseDocParser
        parser = CourseDocParser(docx_paths=[path])

        # Build para_idx → heading level map using A0's proven heuristics.
        heading_map: dict[int, int] = {
            h["para_idx"]: h["level"]
            for h in parser.extract_heading_tree()
        }

        doc = python_docx.Document(path)
        nodes: list[DocumentNode] = []

        for raw_index, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            heading_level = heading_map.get(raw_index)

            if heading_level is not None:
                block_type = BlockType.HEADING
                level = heading_level
            else:
                style_name = para.style.name if para.style else ""
                is_list = style_name.startswith("List")
                block_type = BlockType.LIST_ITEM if is_list else BlockType.PARAGRAPH
                level = 0

            nodes.append(DocumentNode(
                node_id=uuid.uuid4().hex[:12],
                block_type=block_type,
                level=level,
                text=text,
                page_num=None,
                raw_index=raw_index,
            ))

        # Tables: walk body elements to find <w:tbl> blocks and emit TABLE nodes.
        tbl_index = len(doc.paragraphs)
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "tbl":
                try:
                    table = python_docx.table.Table(element, doc)
                    cells = [
                        cell.text.strip()
                        for row in table.rows
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    combined = " | ".join(cells)
                    if combined:
                        nodes.append(DocumentNode(
                            node_id=uuid.uuid4().hex[:12],
                            block_type=BlockType.TABLE,
                            level=0,
                            text=combined,
                            page_num=None,
                            raw_index=tbl_index,
                        ))
                except Exception as exc:
                    logger.warning("[docx_parser] Table parse failed: %s", exc)
                tbl_index += 1

        logger.info("[docx_parser] Parsed %d nodes from %s", len(nodes), path)
        return nodes
