"""Tests for POST /jobs/{job_id}/artifacts/render-docx (render-only, no persistence)."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.api.deps import get_db
from app.api.v1.endpoints.content_generation import course_generation_job as job_routes
from app.core.auth.dependencies import require_valid_token
from app.schemas.onboarding.course_generation_job.course_content_snapshot import (
    RenderDocxRequest,
)
from app.services.onboarding.course_generation.docx_render_service import (
    DocxRenderService,
    sanitize_docx_filename,
)


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _valid_snapshot(**overrides) -> dict:
    payload = {
        "jobId": "42",
        "courseTitle": "Annuity CE Course!",
        "courseType": "Insurance CE",
        "generatedAt": "2026-07-12T10:00:00Z",
        "meta": {
            "totalWordCount": 40,
            "sectionCount": 3,
            "chapterCount": 2,
            "estimatedReadTime": "1 min read",
        },
        "sections": [
            {
                "id": "42-introduction",
                "title": "Introduction",
                "level": 1,
                "sectionType": "overview",
                "content": "Course introduction text.",
                "paragraphs": [
                    {"type": "text", "content": "Course introduction text."}
                ],
                "learningObjectives": [],
                "wordCount": 3,
                "hasKnowledgeCheck": False,
                "order": 0,
                "children": [],
            },
            {
                "id": "ch-1",
                "title": "Chapter One",
                "level": 1,
                "sectionType": "overview",
                "content": "",
                "paragraphs": [],
                "learningObjectives": [],
                "wordCount": 0,
                "hasKnowledgeCheck": False,
                "order": 1,
                "children": [
                    {
                        "id": "sec-manual",
                        "title": "Manual Edit Section",
                        "level": 2,
                        "sectionType": "content",
                        "content": "MANUAL_EDIT_MARKER unique text",
                        "paragraphs": [],
                        "learningObjectives": [],
                        "wordCount": 4,
                        "hasKnowledgeCheck": False,
                        "order": 0,
                        "children": [],
                    },
                    {
                        "id": "sec-ai",
                        "title": "AI Edit Section",
                        "level": 2,
                        "sectionType": "content",
                        "content": "",
                        "paragraphs": [
                            {
                                "type": "text",
                                "content": "AI_EDIT_MARKER regenerated prose",
                            }
                        ],
                        "learningObjectives": [],
                        "wordCount": 4,
                        "hasKnowledgeCheck": False,
                        "order": 1,
                        "children": [],
                    },
                ],
            },
            {
                "id": "ch-2",
                "title": "Added Chapter",
                "level": 1,
                "sectionType": "content",
                "content": "ADDED_SECTION_MARKER brand new",
                "paragraphs": [],
                "learningObjectives": [],
                "wordCount": 4,
                "hasKnowledgeCheck": False,
                "order": 2,
                "children": [],
            },
            {
                "id": "42-conclusion",
                "title": "Conclusion",
                "level": 1,
                "sectionType": "conclusion",
                "content": "CONCLUSION_MARKER wrap up",
                "paragraphs": [
                    {"type": "text", "content": "CONCLUSION_MARKER wrap up"}
                ],
                "learningObjectives": [],
                "wordCount": 3,
                "hasKnowledgeCheck": False,
                "order": 3,
                "children": [],
            },
        ],
    }
    payload.update(overrides)
    return payload


def _docx_plain_text(docx_bytes: bytes) -> str:
    document = Document(io.BytesIO(docx_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _is_zip_docx(docx_bytes: bytes) -> bool:
    bio = io.BytesIO(docx_bytes)
    if not zipfile.is_zipfile(bio):
        return False
    bio.seek(0)
    with zipfile.ZipFile(bio) as zf:
        names = set(zf.namelist())
    return "word/document.xml" in names and "[Content_Types].xml" in names


@pytest.fixture
def render_client():
    app = FastAPI()
    app.include_router(job_routes.router)
    app.dependency_overrides[require_valid_token] = lambda: {"name": "test-user"}
    app.dependency_overrides[get_db] = lambda: MagicMock()

    mock_job = MagicMock()
    mock_job.id = 42

    with patch.object(
        job_routes.CourseGenerationJobRepository,
        "get_by_id",
        return_value=mock_job,
    ) as get_by_id:
        client = TestClient(app)
        yield client, get_by_id

    app.dependency_overrides.clear()


def test_sanitize_docx_filename():
    assert sanitize_docx_filename("Annuity CE Course!") == "Annuity-CE-Course.docx"
    assert sanitize_docx_filename("   ") == "course.docx"


def test_render_service_returns_valid_docx_bytes():
    rendered = DocxRenderService().render(
        RenderDocxRequest.model_validate(_valid_snapshot())
    )
    assert rendered.media_type == DOCX_MEDIA_TYPE
    assert rendered.filename == "Annuity-CE-Course.docx"
    assert _is_zip_docx(rendered.content)
    text = _docx_plain_text(rendered.content)
    assert "MANUAL_EDIT_MARKER" in text
    assert "AI_EDIT_MARKER" in text
    assert "ADDED_SECTION_MARKER" in text
    assert "CONCLUSION_MARKER" in text
    assert "Course introduction text." in text
    assert "OMITTED_SHOULD_NOT_APPEAR" not in text


def test_render_service_respects_chapter_order():
    snapshot = _valid_snapshot()
    snapshot["sections"][1]["order"] = 5  # Chapter One later
    snapshot["sections"][2]["order"] = 1  # Added Chapter earlier
    snapshot["sections"][3]["order"] = 9

    rendered = DocxRenderService().render(RenderDocxRequest.model_validate(snapshot))
    text = _docx_plain_text(rendered.content)
    added_idx = text.index("ADDED_SECTION_MARKER")
    chapter_one_idx = text.index("MANUAL_EDIT_MARKER")
    assert added_idx < chapter_one_idx


def test_render_service_cleans_up_temp_files(monkeypatch):
    created: list[Path] = []
    real_temp_dir = __import__("tempfile").TemporaryDirectory

    class TrackingTemporaryDirectory(real_temp_dir):
        def __enter__(self):
            path = super().__enter__()
            created.append(Path(path))
            return path

    monkeypatch.setattr(
        "app.services.onboarding.course_generation.docx_render_service.TemporaryDirectory",
        TrackingTemporaryDirectory,
    )

    DocxRenderService().render(RenderDocxRequest.model_validate(_valid_snapshot()))
    assert created, "Expected a temporary directory to be created"
    for path in created:
        assert not path.exists(), f"Temp dir was not cleaned up: {path}"


def test_render_docx_endpoint_success(render_client):
    client, get_by_id = render_client

    with (
        patch(
            "app.services.onboarding.course_generation.artifact_service."
            "CourseGenerationArtifactService.persist_bytes"
        ) as persist_bytes,
        patch(
            "app.core.storage.azure_blob_client.AzureBlobClient.upload_bytes"
        ) as upload_bytes,
    ):
        response = client.post(
            "/jobs/42/artifacts/render-docx",
            json=_valid_snapshot(),
        )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(DOCX_MEDIA_TYPE)
    disposition = response.headers.get("content-disposition", "")
    assert "attachment" in disposition
    assert "Annuity-CE-Course.docx" in disposition
    assert _is_zip_docx(response.content)

    text = _docx_plain_text(response.content)
    assert "MANUAL_EDIT_MARKER" in text
    assert "AI_EDIT_MARKER" in text
    assert "ADDED_SECTION_MARKER" in text

    get_by_id.assert_called_once_with("42")
    persist_bytes.assert_not_called()
    upload_bytes.assert_not_called()


def test_render_docx_endpoint_invalid_payload_returns_422(render_client):
    client, _ = render_client
    response = client.post(
        "/jobs/42/artifacts/render-docx",
        json={"courseTitle": "Broken", "sections": "not-a-list"},
    )
    assert response.status_code == 422


def test_render_docx_endpoint_empty_content_returns_422(render_client):
    client, _ = render_client
    response = client.post(
        "/jobs/42/artifacts/render-docx",
        json={"courseTitle": "Empty", "sections": []},
    )
    assert response.status_code == 422
    assert "No renderable course" in response.json()["detail"]


def test_render_docx_endpoint_unknown_job_returns_404(render_client):
    client, get_by_id = render_client
    get_by_id.return_value = None
    response = client.post(
        "/jobs/999/artifacts/render-docx",
        json=_valid_snapshot(),
    )
    assert response.status_code == 404


def test_render_docx_does_not_call_persistence_or_azure():
    """Service-level: rendering never touches artifact persist or Azure clients."""
    with (
        patch(
            "app.services.onboarding.course_generation.artifact_service."
            "CourseGenerationArtifactService.persist_bytes"
        ) as persist_bytes,
        patch(
            "app.services.onboarding.course_generation.artifact_service."
            "CourseGenerationArtifactService.persist_file"
        ) as persist_file,
        patch(
            "app.core.storage.azure_blob_client.AzureBlobClient.upload_bytes"
        ) as upload_bytes,
        patch(
            "app.repositories.course_generation.course_generation_job_repository."
            "CourseGenerationJobRepository.update",
            create=True,
        ) as job_update,
    ):
        DocxRenderService().render(RenderDocxRequest.model_validate(_valid_snapshot()))

    persist_bytes.assert_not_called()
    persist_file.assert_not_called()
    upload_bytes.assert_not_called()
    job_update.assert_not_called()


def test_existing_course_content_mapping_unaffected():
    """Render-docx additions must not alter A2 → editor mapping behaviour."""
    from app.services.onboarding.course_generation.course_content_service import (
        _map_a2_output,
    )

    result = _map_a2_output(
        "1",
        {
            "course_title": "Test Course",
            "course_description": "Intro text.",
            "course_conclusion": "Bye.",
            "sections": [
                {
                    "heading": "Primary purposes",
                    "level": 2,
                    "outline_lesson": "1.0 Why Annuities Matter",
                    "body_paragraphs": [{"type": "text", "content": "body"}],
                    "word_count": 1,
                    "section_id": "",
                    "images": [],
                }
            ],
            "timestamp": "2026-07-12T09:21:22Z",
        },
        course_type="Insurance CE",
    )
    titles = [s["title"] for s in result["sections"]]
    assert titles[0] == "Introduction"
    assert titles[-1] == "Conclusion"
    assert "Why Annuities Matter" in titles[1]
